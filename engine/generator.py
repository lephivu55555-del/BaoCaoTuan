import json
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta

class ReportGenerator:
    def __init__(self, start_date, end_date):
        self.start_date = start_date # str DD/MM/YYYY
        self.end_date = end_date # str DD/MM/YYYY
        self.doc = Document()
        
        # Load config
        config_path = os.path.join(os.path.dirname(__file__), 'config.json')
        with open(config_path, 'r', encoding='utf-8') as f:
            self.config = json.load(f)
            
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(13)

    def generate(self, parsed_data, output_path="bao_cao_tuan.docx"):
        """
        Generates the report based on the provided template and parsed data.
        """
        self._add_header()
        
        # Add daily content
        start_dt = datetime.strptime(self.start_date, "%d/%m/%Y")
        end_dt = datetime.strptime(self.end_date, "%d/%m/%Y")
        
        delta = end_dt - start_dt
        days = [ (start_dt + timedelta(days=i)).strftime("%d/%m/%Y") for i in range(delta.days + 1) ]
        
        for date_str in days:
            events = parsed_data.get(date_str, [])
            self._add_day_content(date_str, events)
            
        self._add_footer()
        
        self.doc.save(output_path)
        return output_path

    def _add_header(self):
        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"BÁO CÁO CÔNG VIỆC TỔ KỸ THUẬT SỐ 02\n(Từ ngày {self.start_date} đến {self.end_date})")
        run.bold = True
        run.font.size = Pt(14)
        
        # Personnel
        p = self.doc.add_paragraph()
        p.add_run("Cán bộ thực hiện: ").bold = True
        p.add_run(self.config['personnel'])
        
        # Section Fixed
        self.doc.add_paragraph().add_run("1. Các công việc đã hoàn thành trong tuần").bold = True
        
        for task in self.config['fixed_tasks']:
            p = self.doc.add_paragraph(task, style='List Bullet')
            p.paragraph_format.line_spacing = 1.15
        
        # Thêm khoảng trắng trước khi in bảng sự cố
        self.doc.add_paragraph()

    def _add_day_content(self, date_str, events):
        p_date = self.doc.add_paragraph()
        p_date.add_run(f"Ngày {date_str}").bold = True
        self.doc.add_paragraph() # Spacer
        
        if not events:
            txt = "Hệ thống giao thông thông minh ITS, hệ thống thu phí ETC, hệ thống truyền dẫn kỹ thuật số DTS, hệ thống kiểm tra tải trọng xe các trạm thu phí đều hoạt động bình thường, không phát sinh hư hỏng."
            p = self.doc.add_paragraph(txt)
            self.doc.add_paragraph() # Spacer
        else:
            for ev in events:
                loc = ev['location']
                self.doc.add_paragraph(f"*{loc}")
                self.doc.add_paragraph() # Spacer
                
                contents = ev['content']
                if isinstance(contents, str):
                    contents = [contents]
                    
                for c in contents:
                    self.doc.add_paragraph(f"- {c}")
                    self.doc.add_paragraph() # Spacer

    def _add_footer(self):
        # Add footers
        self.doc.add_paragraph()
        p1 = self.doc.add_paragraph()
        p1.add_run("2. Các công việc bị chậm tiến độ/tồn tại/vướng mắc: ").bold = True
        p1.add_run("Không.")
        
        self.doc.add_paragraph()
        p2 = self.doc.add_paragraph()
        p2.add_run("3. Kế hoạch tuần tiếp theo:").bold = True
        
        for t in self.config['next_week_plan']:
            self.doc.add_paragraph(f"- {t}")
            
        self.doc.add_paragraph()
        p3 = self.doc.add_paragraph()
        p3.add_run("4. Các đề xuất, kiến nghị, sáng kiến: ").bold = True
        p3.add_run("Không có.")
