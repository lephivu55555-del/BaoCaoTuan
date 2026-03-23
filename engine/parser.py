import os
from docx import Document
import re
from datetime import datetime

class LogParser:
    def __init__(self, file_path):
        self.file_path = file_path
        self.data = {} # {date_str: [list of events]}
        
        self.RE_DATE = re.compile(r"^\s*\d{1,2}[/-]\d{1,2}[/-]\d{4}\s*$")
        self.RE_TIME = re.compile(r"^\s*\d{1,2}[h:]\d{2}\b")
        self.PERIODIC_KEYWORDS = [
            "kiểm tra các thiết bị", "kiểm tra hệ thống", "kiểm tra tổng quan", 
            "kiểm tra phần mềm", "kiểm tra định kỳ"
        ]

    def _cell_text(self, cell) -> str:
        return ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip()).strip()

    def _is_date_row(self, row) -> bool:
        for cell in row.cells:
            if self._cell_text(cell) and re.search(r'(\d{1,2}[/-]\d{1,2}[/-]\d{4})', self._cell_text(cell)):
                return True
        return False

    def _extract_date(self, row):
        for cell in row.cells:
            match = re.search(r'(\d{1,2}[/-]\d{1,2}[/-]\d{4})', self._cell_text(cell))
            if match:
                return match.group(1).replace('-', '/')
        return None

    def _row_has_time(self, row) -> bool:
        for cell in row.cells:
            if self.RE_TIME.match(self._cell_text(cell)):
                return True
        return False

    def _is_periodic_entry(self, texts: list) -> bool:
        combined = " ".join(texts).lower()
        
        # Gỡ bỏ các cụm từ an toàn đánh lừa bộ đếm (mang nghĩa phủ định sự cố)
        safe_phrases = [
            "không phát sinh hư hỏng", "không có hiện tượng hư hỏng", 
            "không phát hiện sự cố", "không có sự cố", "không phát sinh sự cố",
            "chưa phát hiện lỗi", "không có lỗi", "vị trí, lý trình xảy ra phát hiện sự cố"
        ]
        for sp in safe_phrases:
            combined = combined.replace(sp, "")
        
        # Các cụm từ bắt buộc là sự cố hỏng hóc
        incident_markers = [
            "sự cố", "lỗi", "hư hỏng", "không hoạt", "không tự động", 
            "không hạ", "không đọc", "mất kết nối", "đứt", "chập", "cháy", 
            "bị kẹt", "không ổn định", "treo"
        ]
        
        if "hoạt động bình thường" in combined and not any(m in combined for m in incident_markers):
            return True
            
        if any(m in combined for m in incident_markers):
            return False
            
        return any(kw in combined for kw in self.PERIODIC_KEYWORDS)

    def parse(self, start_date_str, end_date_str):
        start_date = datetime.strptime(start_date_str, "%d/%m/%Y")
        end_date = datetime.strptime(end_date_str, "%d/%m/%Y")
        
        doc = Document(self.file_path)
        current_date_obj = None
        current_date_str = None
        
        for table in doc.tables:
            rows = table.rows
            i = 0
            while i < len(rows):
                row = rows[i]
                
                if self._is_date_row(row):
                    raw_date = self._extract_date(row)
                    if raw_date:
                        try:
                            current_date_obj = datetime.strptime(raw_date, "%d/%m/%Y")
                            current_date_str = current_date_obj.strftime("%d/%m/%Y")
                        except:
                            pass
                    i += 1
                    continue
                
                # Check if current date is within bounds
                if not current_date_obj or not (start_date <= current_date_obj <= end_date):
                    i += 1
                    continue
                
                if self._row_has_time(row):
                    texts = ["", "", "", ""]
                    start_i = i
                    
                    in_signature_block = False
                    while i < len(rows):
                        cur_row = rows[i]
                        if i != start_i:
                            if self._row_has_time(cur_row) or self._is_date_row(cur_row):
                                break
                            
                            is_end = any(
                                kw in self._cell_text(c).lower()
                                for c in cur_row.cells[:4]
                                for kw in ("vị trí, lý trình", "nội dung kiểm tra", "giờ ngày", "kết quả")
                            )
                            if is_end:
                                break
                                
                            if not in_signature_block:
                                if any(kw in self._cell_text(c).lower() for c in cur_row.cells[:4] for kw in ("ca trực", "tổ trưởng", "nhân viên", "người lập", "trực ca")):
                                    in_signature_block = True
                                    
                        for col_idx, cell in enumerate(cur_row.cells[:4]):
                            ctext = self._cell_text(cell).strip()
                            if ctext:
                                if in_signature_block and col_idx in (1, 2):
                                    continue
                                texts[col_idx] = (texts[col_idx] + " " + ctext).strip() if texts[col_idx] else ctext
                        i += 1
                        
                    if self._is_periodic_entry(texts):
                        continue
                        
                    if not texts[0].strip():
                        continue
                        
                    if current_date_str not in self.data:
                        self.data[current_date_str] = []
                        
                    vi_tri = texts[1] if texts[1] else "Không rõ vị trí"
                    su_co = texts[2] if texts[2] else ""
                    xu_ly = texts[3] if texts[3] else ""
                    
                    if su_co and xu_ly:
                        noi_dung = f"{su_co}, {xu_ly}"
                    elif su_co:
                        noi_dung = su_co
                    elif xu_ly:
                        noi_dung = xu_ly
                    else:
                        noi_dung = "Không có chi tiết"

                    # Additional cleanup just in case
                    thoi_gian = re.sub(r'\d{1,2}[h:]\d{2}', '', texts[0]).strip()
                    if thoi_gian: 
                        if "hoạt động bình thường" not in noi_dung.lower():
                            noi_dung = f"{noi_dung}"
                        
                    self.data[current_date_str].append({
                        "location": vi_tri,
                        "content": noi_dung
                    })
                    continue
                i += 1
                
        # Group by location to make it clean
        for date_key in self.data:
            merged = []
            loc_map = {}
            for entry in self.data[date_key]:
                loc = entry['location']
                content = entry['content']
                if loc in loc_map:
                    idx = loc_map[loc]
                    if content not in merged[idx]['content']:
                        merged[idx]['content'].append(content)
                else:
                    loc_map[loc] = len(merged)
                    merged.append({"location": loc, "content": [content]})
            self.data[date_key] = merged

        return self.data
